"""Generate the IEEE Access Response-to-Reviewers .docx from the final revision."""
import os
from docx import Document
from docx.shared import Pt, RGBColor

OUT = os.path.join(os.path.dirname(__file__), "..", "..", "EyanaSSDSim-Paper",
                   "Response-to-Reviewers-EyanaSSDSim.docx")

MID = "Access-2026-12395"
TITLE = ("EyanaSSDSim: Explore the Inner Workings of Solid-State Drives with "
         "Data Visualization")

# (reviewer, concern-text, response, action)
ITEMS = [
 ("Reviewer#1, Concern # 1",
  "The simulator is deterministic and does not model stochastic variability such "
  "as hardware latency, bus contention, and bit errors, risking an idealized model.",
  "EyanaSSDSim is deterministic by design, in common with widely used FTL research "
  "simulators (FTLSim, the DiskSim-based SSD model, and default MQSim). Determinism "
  "is deliberate for a visualization-and-analysis tool: it makes every internal event "
  "exactly reproducible, which is what enables step-by-step visual tracking and "
  "repeatable teaching. We now state explicitly that EyanaSSDSim models the logical "
  "and architectural behavior of the FTL (data placement, GC, wear-leveling, WAF), "
  "not analog NAND phenomena (raw bit-error rate, read-retry, program/erase-time "
  "variation) or host/OS overheads.",
  "added a Limitations subsection scoping the model, stated the determinism rationale, "
  "and noted that a stochastic latency/error layer (for which the workload generator "
  "already exposes a seed hook) is future work."),

 ("Reviewer#1, Concern # 2",
  "Only a greedy GC policy is used, which limits the tool for investigating advanced "
  "GC policies (e.g., cost-benefit, windowed).",
  "We rebuilt the simulation core with a pluggable GC interface and implemented two "
  "additional policies -- cost-benefit (age/utilization-weighted) and FIFO -- beyond "
  "greedy. A comparative evaluation shows the cost-benefit policy achieves more uniform "
  "wear than greedy for the uniform random workload (Gini 0.052 -> 0.045), while greedy "
  "attains the lowest write amplification for the Zipf workload; FIFO is the weakest for "
  "Zipf. This demonstrates the engine is not tied to greedy and can serve as a testbed "
  "for reclaim policies.",
  "added a pluggable GC interface with cost-benefit and FIFO policies, together with a "
  "comparison figure (WAF and Gini per policy) and discussion in the Discussion section."),

 ("Reviewer#1, Concern # 3",
  "The tool is limited to a small capacity (a few GB) and its performance can be affected "
  "by browser limits, positioning it as educational rather than a robust research simulator.",
  "The simulation core is headless and configurable; we added a scalability evaluation at "
  "5, 20, and 40 GB, sustaining near-constant throughput (about 4.3-4.5x10^5 write "
  "operations per second). Only the browser-based visualization layer is bounded by "
  "rendering, which is addressed by the aggregated view (below). The device size in the "
  "paper was chosen for figure legibility, not as an engine limit.",
  "added a scalability table (5/20/40 GB with throughput) and clarified that capacity is "
  "configurable and independent of the visualization layer."),

 ("Reviewer#1, Concern # 4",
  "The proposed 'aggregated views' for scalability are not described precisely, risking loss "
  "of the real-time observation that is the main contribution.",
  "We now describe the aggregation precisely: at large capacities each rendered cell maps to "
  "a fixed group of physical blocks and its color encodes the group's mean invalid-page ratio "
  "(and, on toggle, mean erase count). This preserves observation of spatial trends in real "
  "time while bounding the number of drawn elements.",
  "added a description of the aggregation scheme to the architecture/discussion."),

 ("Reviewer#1, Concern # 5",
  "The WAF differences versus FTLSim are attributed vaguely to 'different simulation approaches' "
  "and to LSB backup pages, suggesting an oversimplified model.",
  "During the revision we rebuilt the engine and, importantly, discovered that our earlier "
  "FTLSim/FEMU comparison had used a larger device than FTLSim/FEMU. Re-running EyanaSSDSim on "
  "the same 1.27 GB device used for FTLSim/FEMU, the synthetic workloads now agree closely: at "
  "10% OPS EyanaSSDSim reports WAF 2.39 (uniform) and 1.13 (Zipf) versus 2.55 and 1.13 for "
  "FTLSim. The remaining sequential/prxy gaps are quantitatively explained by FTLSim's LSB "
  "backup pages (e.g., 2.8M backup pages for prxy inflate its WAF), not by our WAF calculation. "
  "We now state this explicitly and bound the difference.",
  "re-ran the validation on the matched device size, updated the WAF comparison figure, and "
  "added text quantifying the close synthetic-workload agreement and the LSB-backup explanation."),

 ("Reviewer#1, Concern # 6",
  "The physical justification for using Fourier analysis to measure wear-leveling, versus "
  "standard statistical dispersion metrics (e.g., coefficient of variation, Gini index), is unclear.",
  "We now report the Coefficient of Variation and the Gini coefficient of the per-block erase "
  "counts alongside the Fourier amplitude spread. The scalar metrics quantify how uneven the wear "
  "is, while the Fourier spectrum additionally exposes the spatial/periodic structure of erase "
  "concentration (which groups of blocks are repeatedly hottest, and at what spatial frequency) "
  "that a single aggregate number cannot. We present Fourier as complementary to, not a "
  "replacement for, CV and Gini. We also removed the DC component from the amplitude spread so "
  "that a perfectly uniform distribution correctly yields zero spread.",
  "added CV and Gini columns/metrics and a paragraph motivating the Fourier analysis as "
  "complementary to standard dispersion measures."),

 ("Reviewer#1, Concern # 7",
  "Verify that the Evaluation section defines the DoIPD and DoEC metrics in sufficient detail so "
  "the reader need not jump between sections.",
  "Both metrics are defined with their equations, and a summary notation table together with an "
  "'Evaluation Metrics' list (WAF, DoIPD, DoEC, CV, Gini) is provided in the evaluation, so the "
  "reader does not need to jump between sections.",
  "confirmed that the DoIPD and DoEC definitions, their equations, and a notation table are present "
  "in the Evaluation section."),

 ("Reviewer#1, Concern # 8 (additional)",
  "The GC thresholds (start at 99.9995%, stop at 99.9990%) are extremely narrow and unrealistic, "
  "and could bias the WAF results.",
  "This was an artifact of how the earlier prototype expressed its trigger. The rebuilt engine uses "
  "standard high/low block-utilization watermarks (start GC at 95%, stop at 90%), which reflect "
  "typical SSD over-provisioning behavior. All reported results use these watermarks.",
  "replaced the near-full trigger with 95%/90% watermarks in the device table and text."),

 ("Reviewer#1, Concern # 9 (additional)",
  "It is not specified whether the reported read latency accounts for bus contention or OS overhead, "
  "weakening the validation.",
  "We clarify that the reported latency models per-operation flash service time (page read/program, "
  "block erase) with channel-level serialization, not host-stack or operating-system overhead, and we "
  "scope the FEMU latency comparison accordingly.",
  "clarified the latency model's scope in the validation section."),

 ("Reviewer#1, Concern # 10 (additional)",
  "Dominant/emerging technologies such as QLC/PLC and ZNS/FDP, and recent (2024-2026) work, are only "
  "mentioned superficially.",
  "We expanded the Discussion to cite recent Zoned Namespace work (ZNS and ZNS+), Flexible Data "
  "Placement work (the FDP open-source ecosystem and FlashAlloc), and Quadruple-Level Cell (QLC) "
  "NAND analysis. We also explain why the quantitative wear-leveling analysis presented here (DoIPD, "
  "DoEC, CV, Gini, and Fourier spread) is increasingly relevant for high-density QLC and PLC NAND, "
  "which are more sensitive to write amplification and uneven wear. As a direction for future work, "
  "we describe a zone-aware extension of EyanaSSDSim with host-managed garbage collection and "
  "hot/cold zone separation for ZNS-specific analysis.",
  "expanded the Discussion with recent ZNS, FDP, and QLC references and outlined the ZNS extension."),

 ("Reviewer#1, Concern # 11 (additional)",
  "Internal contradiction: Section IV-B states the uniform random pattern is beneficial for "
  "wear-leveling, while the erase-count analysis concludes the sequential pattern is optimal.",
  "This was a genuine ambiguity that we have resolved. The apparent contradiction arose from "
  "conflating two distinct notions: total wear (mean erase count, minimized by low-WAF workloads) "
  "and wear uniformity (dispersion of erase counts). We now state the trade-off explicitly and "
  "consistently: the uniform random workload yields the most uniform wear (lowest Gini/CV) but "
  "higher total wear; the sequential workload incurs the least total wear (a consequence of its "
  "near-ideal WAF); and the Zipf workload is the least uniform. All four metrics (DoEC, CV, Gini, "
  "Fourier) are reported together so the reader sees one coherent story.",
  "rewrote Section IV-B and the erase-count analysis to state the two-axis trade-off up front and "
  "removed the contradictory phrasing."),

 ("Reviewer#1, Concern # 12 (additional)",
  "A satisfaction survey (95% of 1,011 participants) is a usability metric and does not validate the "
  "technical accuracy of the simulator.",
  "We agree and now scope the survey explicitly as usability/educational evidence, separate from the "
  "accuracy validation (which is established against FTLSim and FEMU). Technical accuracy and usability "
  "are presented as distinct claims.",
  "reframed the Survey section as a usability/educational result and separated it from the accuracy "
  "validation."),

 ("Reviewer#2, Concern # 1",
  "The internal logic (e.g., a greedy GC policy) may be too simplistic compared to proprietary FTLs "
  "in industrial SSDs.",
  "We rebuilt the simulation core with a pluggable garbage-collection interface and now support "
  "cost-benefit and FIFO policies in addition to greedy. A comparison (added as a figure) shows the "
  "cost-benefit policy achieves more uniform wear than greedy on the uniform random workload, while "
  "greedy attains the lowest write amplification on the Zipf workload. We also scope the tool's claims "
  "to the FTL layer it models (data placement, garbage collection, wear-leveling, and write "
  "amplification); we acknowledge that proprietary industrial FTLs are more complex, and position "
  "EyanaSSDSim as an open, extensible analytical and educational testbed rather than a replica of a "
  "specific commercial device.",
  "added cost-benefit and FIFO GC policies with a comparison figure, and scoped the modeling claims "
  "to the FTL layer."),

 ("Reviewer#2, Concern # 2",
  "The results are largely descriptive of what is happening visually, rather than providing a rigorous "
  "proof of predictive accuracy.",
  "Beyond visualization we provide quantitative, reproducible metrics (WAF, DoIPD, DoEC, CV, Gini, "
  "Fourier) and validate WAF and latency against FTLSim and FEMU. The Workload Similarity Composition "
  "(WSC) metric maps real traces to synthetic profiles, enabling prescriptive firmware-tuning guidance "
  "(e.g., OPS sizing), which we make explicit.",
  "strengthened the analysis framing and emphasized the WSC-driven, prescriptive guidance."),

 ("Reviewer#2, Concern # 3",
  "The work lacks a deep performance/accuracy benchmark against established tools.",
  "We report a direct WAF comparison against FTLSim (10% OPS) and FEMU (25% OPS) across all workloads "
  "and a read-latency comparison against FEMU, with differences quantified and explained. The engine "
  "now has a regression test suite that pins its numerical behavior.",
  "expanded the validation and added a unit-tested engine."),

 ("Reviewer#2, Concern # 4",
  "The Survey Result section is shallow: it does not present the survey or discuss the results, and a "
  "controlled experiment would be needed to evaluate educational usefulness.",
  "We reframed the survey as usability/educational evidence and now present its scope and the "
  "distribution of responses. We agree that a controlled learning study would further strengthen the "
  "educational claim, and we note it as planned future work.",
  "expanded and reframed the Survey section and scoped its claims."),

 ("Reviewer#2, Concern # 5",
  "The experimental design lacks a formal baseline comparison and evidence that results match real "
  "hardware.",
  "We provide the FTLSim/FEMU baseline comparison (WAF and latency) and clarify that EyanaSSDSim models "
  "FTL-level behavior, validated against these established simulators rather than against a specific "
  "hardware device. The manuscript now also states that the rebuilt engine is deterministic and "
  "unit-tested, with a regression suite pinning the WAF, GC, and allocation computations for "
  "reproducibility.",
  "added the baseline comparison, clarified the validation target, and noted that the engine is "
  "deterministic and unit-tested for reproducibility."),

 ("Reviewer#3, Concern # 1",
  "Clarify the novelty: what is fundamentally new beyond visualization, and how it differs "
  "quantitatively from existing simulators.",
  "We sharpened the contribution statement. The novelty is the combination of (i) real-time internal "
  "visualization, (ii) the WSC metric that maps real traces to synthetic profiles for workload-aware "
  "tuning, and (iii) a multi-metric wear analysis (DoEC/CV/Gini/Fourier) not provided by the surveyed "
  "simulators; the feature table quantifies the gap versus MQSim, SSDModel, FEMU, and others.",
  "rewrote the contributions and the differentiation discussion/feature table."),

 ("Reviewer#3, Concern # 2",
  "Improve technical depth by supporting/discussing more advanced GC strategies and queueing models, "
  "and justify why deterministic simulation is sufficient.",
  "We improved the technical depth in two ways. First, we added two garbage-collection policies beyond "
  "greedy -- cost-benefit (age/utilization-weighted) and FIFO -- with a comparative evaluation of write "
  "amplification and wear across the workloads. Second, we justify the deterministic design: "
  "EyanaSSDSim is deterministic by choice, as are established FTL research simulators, because "
  "determinism makes every internal event exactly reproducible, which is precisely what enables "
  "step-by-step visual tracking and repeatable teaching. We added an explicit Limitations statement "
  "noting that stochastic latency/error effects and advanced queueing models are outside the current "
  "scope and are planned future work.",
  "added cost-benefit and FIFO GC policies with a comparison, and added a determinism rationale and a "
  "Limitations statement."),

 ("Reviewer#3, Concern # 3",
  "The Fourier Transform analysis feels over-engineered relative to the insight gained and needs "
  "clearer motivation versus simpler statistical methods.",
  "We now report the Coefficient of Variation and the Gini coefficient of the per-block erase counts "
  "alongside the Fourier amplitude spread, and we motivate the Fourier analysis as a complementary "
  "diagnostic. The scalar metrics quantify how uneven the wear is, whereas the Fourier spectrum "
  "additionally exposes the spatial and periodic structure of erase concentration -- which groups of "
  "blocks are repeatedly hottest, and at what spatial frequency -- that a single aggregate value "
  "cannot. We also removed the DC component from the spread so that a perfectly uniform distribution "
  "correctly yields zero. Because the three views agree on the overall ranking, the frequency-domain "
  "analysis is validated by, rather than a replacement for, the standard dispersion metrics.",
  "added CV and Gini metrics alongside the Fourier spread and a clearer, complementary motivation for "
  "the Fourier analysis."),

 ("Reviewer#3, Concern # 4",
  "Fix grammar inconsistencies, improve sentence structure, and reduce overly long explanations.",
  "We fixed specific grammatical errors (e.g., 'invalidation's' -> 'invalidations', 'shwon' -> "
  "'shown', a hard-coded table reference, and a block-count inconsistency) and tightened several "
  "long passages to improve readability.",
  "fixed the specific grammatical errors noted above and tightened several passages."),
]


# where each change can be found in the manuscript (compiled numbering)
LOCATIONS = {
 "Reviewer#1, Concern # 1": "These changes appear in the Limitations paragraph of Section VII and the determinism statement in Section IV.",
 "Reviewer#1, Concern # 2": "See Section VII (Discussion) and the new Figure 13.",
 "Reviewer#1, Concern # 3": "See the new Table 9 (scalability) and Section VII.",
 "Reviewer#1, Concern # 4": "See Section VII (Discussion).",
 "Reviewer#1, Concern # 5": "See Section VI (Performance Validation) and the revised Figure 12.",
 "Reviewer#1, Concern # 6": "See Section V (Analysis of the Simulation Results), Figure 7, and the erase-count analysis in Section V.",
 "Reviewer#1, Concern # 7": "See the Evaluation Metrics list in Section IV, the notation table (Table 6), and the DoIPD/DoEC equations in Section V.",
 "Reviewer#1, Concern # 8 (additional)": "See Table 2 (device specification) and Section III.",
 "Reviewer#1, Concern # 9 (additional)": "See Section VI (Performance Validation).",
 "Reviewer#1, Concern # 10 (additional)": "See Section VII (Discussion and Implication).",
 "Reviewer#1, Concern # 11 (additional)": "See the erase-count analysis in Section V and Figures 5 and 6.",
 "Reviewer#1, Concern # 12 (additional)": "See Section VIII (Survey Result).",
 "Reviewer#2, Concern # 1": "See Section VII (Discussion) and the new Figure 13.",
 "Reviewer#2, Concern # 2": "See the Workload Characterization and Mapping analysis in Section V and Table 7.",
 "Reviewer#2, Concern # 3": "See Section VI (Performance Validation) and Figures 11 and 12.",
 "Reviewer#2, Concern # 4": "See Section VIII (Survey Result).",
 "Reviewer#2, Concern # 5": "See Section VI (Performance Validation), the implementation note in Section IV, and the Code and Data Availability statement (the engine and experiment scripts are publicly released).",
 "Reviewer#3, Concern # 1": "See the contributions list in Section I and the feature comparison in Table 1.",
 "Reviewer#3, Concern # 2": "See Section VII and Figure 13; the determinism rationale is in Section IV and the Limitations paragraph of Section VII.",
 "Reviewer#3, Concern # 3": "See Section V (Analysis of the Simulation Results) and Figure 7.",
 "Reviewer#3, Concern # 4": "These corrections are applied throughout the manuscript.",
}


def _para(doc, label, text, bold_label=True, color=None):
    p = doc.add_paragraph()
    if label:
        r = p.add_run(label + " ")
        r.bold = bold_label
    r2 = p.add_run(text)
    if color:
        r2.font.color.rgb = color
    return p


def main():
    doc = Document()
    for s in doc.styles:
        pass
    doc.add_paragraph(f"Original Manuscript ID: {MID}")
    doc.add_paragraph(f"Original Article Title: “{TITLE}”")
    doc.add_paragraph("")
    doc.add_paragraph("To: IEEE Access Editor")
    doc.add_paragraph("Re: Response to reviewers")
    doc.add_paragraph("")
    doc.add_paragraph("Dear Editor,")
    doc.add_paragraph(
        "Thank you for allowing a resubmission of our manuscript, with an "
        "opportunity to address the reviewers’ comments. We are uploading "
        "(a) our point-by-point response to the comments (below), (b) an updated "
        "manuscript with yellow highlighting indicating changes (“Highlighted "
        "PDF”), and (c) a clean updated manuscript without highlights "
        "(“Main Manuscript”).")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("Summary of the major revision: ").bold = True
    p.add_run(
        "We rebuilt the simulation engine as a headless, unit-tested core. In doing "
        "so we identified and fixed correctness defects in the earlier prototype "
        "(separate counting of host vs. GC writes, a page-per-block off-by-one, and a "
        "configuration under which some allocation schemes were not distinct) and "
        "re-ran all synthetic experiments on the corrected engine. We also matched the "
        "device size to the FTLSim/FEMU comparison, added cost-benefit and FIFO GC "
        "policies, added Coefficient-of-Variation and Gini wear metrics, added a "
        "scalability study, resolved the wear-leveling ambiguity, and scoped the "
        "determinism and survey claims. Where results changed, we report the new, "
        "verified values.")
    doc.add_paragraph("")
    doc.add_paragraph("Best regards,")
    doc.add_paragraph("Md Habibur Rahman et al.")
    doc.add_paragraph("")

    for label, concern, response, action in ITEMS:
        _para(doc, f"{label} (please list here):", concern)
        _para(doc, "Author response:", response)
        loc = LOCATIONS.get(label, "")
        action_text = "We updated the manuscript by " + action
        if loc:
            action_text += " " + loc
        _para(doc, "Author action:", action_text)
        doc.add_paragraph("")

    doc.add_paragraph(
        "Note: References suggested by reviewers were added only where relevant to "
        "the article and where they make it more complete.")
    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    main()
